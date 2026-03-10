import os
import json
import asyncio
import dashscope
from dashscope import MultiModalConversation
from llama_index.core import Document, VectorStoreIndex, StorageContext, Settings,SimpleDirectoryReader
from llama_index.vector_stores.elasticsearch import ElasticsearchStore
from docx import Document as DocxDocument  
from docx.shared import Pt, Inches
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from app.core.kb_manager import UPLOAD_DIR

async def extract_knowledge_from_video(video_filepath: str, video_title: str) -> bool:
    """
    使用阿里原生 SDK 解析超长视频
    """
    print("\n" + "="*50)
    print("🚀 启动原生 DashScope 视频解析模式")
    print("="*50)
    
    try:
        # 1. 自动配置密钥
        dashscope.api_key = os.environ.get("DASHSCOPE_API_KEY")
        if not dashscope.api_key:
            print("❌ 错误：找不到 DASHSCOPE_API_KEY")
            return False

        # 2. 获取本地文件的绝对路径，原生SDK专用格式！
        # 只需要 file:// + 电脑里的绝对路径即可，无需手动上传！
        abs_path = os.path.abspath(video_filepath)
        local_file_url = f"file://{abs_path}"
        
        print(f"🎬 1. 准备读取本地视频: {local_file_url}")
        
        prompt_text = (
            f"你是一个专业的工业设备分析专家。请仔细观看这段名为《{video_title}》的工人实操视频。\n"
            "请将视频中的操作过程拆解为连续的、结构化的操作步骤。\n"
            "警告：绝对不要在 details 中生成或编造任何图片链接、Markdown图片语法（如 ![image](/images/xxx.png)）。只输出纯文本说明。\n"
            "要求严格按照JSON数组格式返回，不带```json标记。\n"
            "格式示例：[{\"step_number\": 1, \"action_name\": \"动作\", \"details\": \"说明\"}]"
        )

        # 3. 构造原生消息结构 (注意这里键名直接叫 'video')
        messages = [
            {
                "role": "user",
                "content": [
                    {"video": local_file_url},
                    {"text": prompt_text}
                ]
            }
        ]

        print("⏳ 2. 正在调用千问大模型 (原生SDK会自动在后台秒传文件)...")
        
        # 4. 由于原生 call 是同步阻塞的，我们把它丢进异步线程池里，防止卡死服务器
        response = await asyncio.to_thread(
            MultiModalConversation.call,
            model='qwen3-vl-plus',  # 这里指定模型
            messages=messages
        )
        
        # 5. 处理结果
        if response.status_code == 200:
            print("✅ 3. 模型解析成功！")
            
            # 提取文本内容
            content = response.output.choices[0].message.content[0].get('text', '')
            print(f"📄 模型原始返回:\n{content}\n")
            
            # 清理可能存在的 Markdown 标记
            content = content.strip()
            if content.startswith("```json"):
                content = content[7:-3]
            elif content.startswith("```"):
                content = content[3:-3]
                
            extracted_steps = json.loads(content)
            print(f"🎉 4. 成功提取了 {len(extracted_steps)} 个结构化步骤！")
            
            # 过滤掉文件名中的非法字符，生成统一的文件名
            safe_title = video_title.replace("/", "_").replace("\\", "_").replace(" ", "")
            word_filename = f"{safe_title}_SOP自动生成.docx"
            word_filepath = os.path.join(UPLOAD_DIR, word_filename)

            # =======================================================
            # 1. 存入 Elasticsearch 数据库
            # =======================================================
            documents = []
            for step in extracted_steps:
                doc_text = f"操作指导来源视频：《{video_title}》。第 {step.get('step_number')} 步：{step.get('action_name')}。详细操作说明：{step.get('details')}"
                doc = Document(
                    text=doc_text,
                    metadata={
                        "source": "video_extraction",
                        "video_title": video_title,
                        "step_number": step.get("step_number"),
                        "action_name": step.get("action_name"),
                        "file_name": word_filename 
                    }
                )
                documents.append(doc)

            es_url = os.environ.get("ES_URL", "http://localhost:9200")
            vector_store = ElasticsearchStore(es_url=es_url, index_name="factory_knowledge")
            storage_context = StorageContext.from_defaults(vector_store=vector_store)
            VectorStoreIndex.from_documents(documents, storage_context=storage_context)
            print(f"✅ 5. 视频知识已存入 ES，归属于文件: {word_filename}")

            # =======================================================
            # 2. 生成 Word 实体文档并保存到 UPLOAD_DIR
            # =======================================================
            from docx import Document as DocxDocument
            from docx.shared import Pt, Inches
            from docx.enum.text import WD_PARAGRAPH_ALIGNMENT

            # 确保目录存在
            os.makedirs(UPLOAD_DIR, exist_ok=True)

            doc_word = DocxDocument()
            title_paragraph = doc_word.add_heading(f"《{video_title}》- 标准操作指导 (SOP)", level=0)
            title_paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

            doc_word.add_paragraph(f"本文档由工厂智能助手根据操作视频自动解析生成。")
            doc_word.add_paragraph(f"视频名称：{video_title}")
            doc_word.add_paragraph(f"生成时间：{__import__('datetime').datetime.now().strftime('%Y-%m-%d %H:%M:%S')}")
            doc_word.add_page_break()

            for step in extracted_steps:
                step_num = step.get("step_number", "")
                action = step.get("action_name", "未知动作")
                details = step.get("details", "无详细说明")

                doc_word.add_heading(f"步骤 {step_num}：{action}", level=1)
                p = doc_word.add_paragraph()
                p.add_run("操作详情：").bold = True
                p.add_run(details)
                doc_word.add_paragraph("")

            # 物理保存 Word 文件
            doc_word.save(word_filepath)
            print(f"📝 6. Word 文档物理文件已保存至: {word_filepath}")
            # =======================================================
            
            return True
        else:
            print(f"❌ 模型调用失败: HTTP {response.status_code}")
            print(f"错误代码: {response.code}")
            print(f"错误信息: {response.message}")
            return False

    except Exception as e:
        print("\n❌❌❌ 捕获到严重异常 ❌❌❌")
        import traceback
        traceback.print_exc()
        return False